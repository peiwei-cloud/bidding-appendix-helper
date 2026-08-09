# -*- coding: utf-8 -*-
"""
備標人員附錄整理、證照檢核與 CV 結構化轉檔系統（100% 本地離線版）
==================================================
功能：
1. 依人員名單 Excel 排序，將每位同仁的 CV / 學歷 / 證照 / 技師執業執照 /
   技師會員證 / 投保證明，依標準 6 分類順序合併成單一附錄 PDF。
2. 逐頁掃描技師執業執照與公會會員證內文字，判讀民國年效期，產出「已過期 /
   即將過期 / 有效」檢核報告。
3. 以純 Python 正則／規則式邏輯解析每位同仁 CV 全文，離線萃取並產出符合
   Template_Staffing.xlsx 標準格式的 14 個欄位，內容已盡量豐富完整，
   方便使用者後續自行複製到 Gemini 網頁版或其他工具做最終文字潤飾。
4. 證照/資格（Licenses／Badges）不套用任何自動縮寫字典：系統僅離線擷取
   每位同仁的證照/資格全稱候選清單，實際是否採用與縮寫代碼，由使用者於
   介面上逐項勾選並手動輸入指定。

⚠️ 本工具 100% 本地離線運作，不呼叫任何外部 AI／雲端 API，不會有任何
   人員資料或標案內容外流，也沒有任何 API 金鑰或費用問題。

部署需求（Streamlit Community Cloud）：
- 本機 / 雲端皆需安裝 LibreOffice（將 .doc/.docx 轉為 PDF 供合併，以及供
  文字擷取備援）。若部署到 Streamlit Cloud，請在 repo 根目錄新增
  `packages.txt`，內容如下（本工具已一併產生於輸出資料夾）：
      libreoffice
      tesseract-ocr
      tesseract-ocr-chi-tra
      poppler-utils

本機執行：
    pip install -r requirements.txt
    streamlit run app.py
"""

import difflib
import io
import os
import re
import subprocess
import tempfile
import zipfile
from datetime import date
from pathlib import Path

import pandas as pd
import streamlit as st
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
import openpyxl

try:
    from docx import Document as DocxDocument
except Exception:  # pragma: no cover
    DocxDocument = None

try:
    from pypdf import PdfReader, PdfWriter
except Exception:  # pragma: no cover
    PdfReader = None
    PdfWriter = None

try:
    import pdfplumber
except Exception:  # pragma: no cover
    pdfplumber = None

try:
    from PIL import Image
except Exception:  # pragma: no cover
    Image = None


# ============================================================
# 常數設定
# ============================================================

CATEGORY_ORDER = [
    "1_CV",
    "2_學歷",
    "3_證照",
    "4_技師執業執照",
    "5_技師會員證",
    "6_投保證明",
]

CATEGORY_LABELS = {
    "1_CV": "CV／履歷",
    "2_學歷": "學歷證明",
    "3_證照": "一般證照",
    "4_技師執業執照": "技師執業執照",
    "5_技師會員證": "技師公會會員證",
    "6_投保證明": "投保證明",
}

CATEGORY_KEYWORDS = {
    "6_投保證明": ["投保", "被保險人", "被保险人", "勞保", "劳保", "勞退"],
    "2_學歷": ["畢業", "毕业", "學歷", "学历", "學位", "学位"],
}

# CV 檔名關鍵字（擴充：學經歷／經歷表）
CV_FILENAME_KEYWORDS = ["cv", "履歷", "履历", "學經歷", "学经历", "經歷表", "经历表"]

# 三段式雙軌配對機制信心度門檻
HIGH_CONFIDENCE_THRESHOLD = 0.80
MEDIUM_CONFIDENCE_THRESHOLD = 0.60

TEMPLATE_COLUMNS = [
    "Layer", "Role", "GroupName", "Name", "Title", "Company", "Licenses", "Badges",
    "PhotoName", "YearsOfExp", "Degree", "JobDescription", "Expertise",
    "BioNarrative",
]

LAYER_OPTIONS = ["Top", "SubTop", "Advisor", "Middle", "GroupLeader",
                  "GroupMember", "Subcontractor"]

EXCLUDE_DIR_TOKENS = ["目前不需要", "不需要", "__macosx", ".git"]


# ============================================================
# 工具函式：檔案分類與人名比對
# ============================================================

def classify_category_by_filename_hint(fname_lower: str) -> str:
    """僅用於投保證明／學歷證明（內容結構化程度低，維持以檔名關鍵字判斷）。
    找不到關鍵字則回傳 None，交由主分類流程繼續判斷。"""
    for cat in ["6_投保證明", "2_學歷"]:
        for kw in CATEGORY_KEYWORDS[cat]:
            if kw.lower() in fname_lower:
                return cat
    return None


# ------------------------------------------------------------
# 三段式雙軌配對機制（姓名 + 英文姓名，容錯拼寫誤差）
# ------------------------------------------------------------

def normalize_name(s: str) -> str:
    """轉小寫並移除常見分隔符號，供比對使用。"""
    s = str(s or "").lower()
    s = re.sub(r"[_\-()（）\s,，、.。]", "", s)
    return s


def _windowed_ratio(alias_norm: str, stem_norm: str) -> float:
    """計算 alias 與檔名（正規化後）的相似度，錨定於檔名最左側區段，
    並輔以全字串比對，兼顧「姓名在最左側」與「姓名夾雜於中段」兩種情況。"""
    if not alias_norm or not stem_norm:
        return 0.0
    if alias_norm in stem_norm:
        return 1.0

    best = 0.0
    lo = max(1, len(alias_norm) - 3)
    hi = min(len(stem_norm), len(alias_norm) + 6)
    for end in range(lo, hi + 1):
        window = stem_norm[:end]
        ratio = difflib.SequenceMatcher(None, alias_norm, window).ratio()
        if ratio > best:
            best = ratio

    full_ratio = difflib.SequenceMatcher(None, alias_norm, stem_norm).ratio()
    return max(best, full_ratio)


def build_alias_list(df_people: pd.DataFrame) -> list:
    """回傳 [(alias_norm, alias_original, person_name), ...]，
    姓名與英文姓名皆納入比對清單。"""
    aliases = []
    for _, row in df_people.iterrows():
        person = str(row["姓名"]).strip()
        for col in ["姓名", "英文姓名"]:
            val = row.get(col, "") if col in df_people.columns else ""
            if val is None or (isinstance(val, float) and pd.isna(val)):
                continue
            val = str(val).strip()
            if val and val.lower() != "nan":
                aliases.append((normalize_name(val), val, person))
    return aliases


def match_person_candidates(filename: str, aliases: list) -> list:
    """回傳依信心分數（0~1）由高到低排序的
    [(person_name, score, matched_alias_text), ...]。"""
    stem = Path(filename).stem
    stem_norm = normalize_name(stem)

    best_per_person = {}
    for alias_norm, alias_text, person in aliases:
        score = _windowed_ratio(alias_norm, stem_norm)
        if person not in best_per_person or score > best_per_person[person][0]:
            best_per_person[person] = (score, alias_text)

    ranked = sorted(
        [(person, score, alias_text) for person, (score, alias_text) in best_per_person.items()],
        key=lambda x: x[1],
        reverse=True,
    )
    return ranked


# ============================================================
# 工具函式：格式轉換
# ============================================================

def convert_to_pdf(filepath: str, workdir: str):
    """將 docx/doc/圖片轉為 PDF；已是 PDF 則直接回傳原路徑。失敗回傳 None。"""
    ext = Path(filepath).suffix.lower()
    if ext == ".pdf":
        return filepath

    if ext in (".doc", ".docx"):
        try:
            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf",
                 "--outdir", workdir, filepath],
                check=True, timeout=180, capture_output=True,
            )
            out_path = Path(workdir) / (Path(filepath).stem + ".pdf")
            return str(out_path) if out_path.exists() else None
        except Exception:
            return None

    if ext in (".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif") and Image is not None:
        try:
            img = Image.open(filepath).convert("RGB")
            out_path = Path(workdir) / (Path(filepath).stem + ".pdf")
            img.save(out_path, "PDF")
            return str(out_path)
        except Exception:
            return None

    return None


def extract_docx_text(filepath: str, workdir: str) -> str:
    """讀取 Word CV 全文（含表格）。.doc 會先透過 LibreOffice 轉為 .docx。"""
    ext = Path(filepath).suffix.lower()
    target = filepath

    if ext == ".doc":
        try:
            subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "docx",
                 "--outdir", workdir, filepath],
                check=True, timeout=180, capture_output=True,
            )
            converted = Path(workdir) / (Path(filepath).stem + ".docx")
            if converted.exists():
                target = str(converted)
        except Exception:
            pass

    if DocxDocument is None:
        return ""

    try:
        doc = DocxDocument(target)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        return "\n".join(parts)
    except Exception:
        return ""


def extract_cv_text(filepath: str, workdir: str) -> str:
    """依副檔名分派 CV 文字擷取方式：.pdf 走 PDF 文字/OCR 擷取，
    .doc/.docx 走 Word 段落/表格擷取。確保 PDF 格式 CV 不再被跳過。"""
    ext = Path(filepath).suffix.lower()
    if ext == ".pdf":
        return extract_pdf_text(filepath)
    if ext in (".doc", ".docx"):
        return extract_docx_text(filepath, workdir)
    return ""


def extract_pdf_pages_text(pdf_path: str) -> list:
    """回傳每一頁的文字內容（list[str]），供逐頁證照效期檢核使用，可正確處理
    「多張證照掃描檔合併於同一份 PDF」（如潘冠愷合併.pdf）的情況。先嘗試
    pdfplumber 文字層，任何頁面缺文字層則對該頁單獨執行 OCR 補強。"""
    pages = []
    if pdfplumber is not None:
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text() or ""
                    pages.append(t)
        except Exception:
            pages = []

    needs_ocr = (not pages) or any(not p.strip() for p in pages)
    if needs_ocr:
        try:
            from pdf2image import convert_from_path
            import pytesseract
            images = convert_from_path(pdf_path, dpi=200)
            if len(pages) != len(images):
                pages = ["" for _ in images]
            for i, img in enumerate(images):
                if not pages[i].strip():
                    pages[i] = pytesseract.image_to_string(img, lang="chi_tra+eng")
        except Exception:
            pass

    return pages


def extract_pdf_text(pdf_path: str) -> str:
    """從 PDF 擷取「整份文件」串接文字（供分類/CV 全文擷取使用），
    若為掃描檔（無文字層）則逐頁嘗試 OCR 備援。"""
    pages = extract_pdf_pages_text(pdf_path)
    return "\n".join(pages)


def find_license_pages(filepath: str, workdir: str) -> list:
    """逐頁掃描 PDF（或先轉檔的圖片/Word），獨立判定每一頁是否為技師執業執照
    或技師公會會員證，並各自進行效期檢核。用於正確處理「多張證照掃描檔合併於
    同一份 PDF」的情況（例如：潘冠愷合併.pdf 內含執業執照頁與會員證頁），
    避免整份檔案只能被歸類為單一種類而漏掉其他頁面的證照。

    回傳 [{"category": "4_技師執業執照"/"5_技師會員證", "check": {...}}, ...]，
    一份檔案若有多頁符合，會產生多筆結果。"""
    ext = Path(filepath).suffix.lower()
    pdf_path = filepath if ext == ".pdf" else convert_to_pdf(filepath, workdir)
    if not pdf_path or not os.path.exists(pdf_path):
        return []

    try:
        pages_text = extract_pdf_pages_text(pdf_path)
    except Exception:
        return []

    results = []
    for page_text in pages_text:
        if not page_text or not page_text.strip():
            continue
        if "技師執業執照" in page_text or "執業執照" in page_text or "执业执照" in page_text:
            check = parse_license_expiry("4_技師執業執照", page_text)
            results.append({"category": "4_技師執業執照", "check": check})
            continue
        has_member = ("會員證" in page_text) or ("会员证" in page_text)
        has_guild = ("技師公會" in page_text) or ("技师公会" in page_text)
        if has_member and has_guild:
            check = parse_license_expiry("5_技師會員證", page_text)
            results.append({"category": "5_技師會員證", "check": check})

    return results


def classify_file_content(filepath: str, workdir: str):
    """內文導向（Content-Driven）分類。回傳 (category, extracted_text)。

    - 1_CV：以檔名關鍵字（CV/履歷/學經歷/經歷表，不分大小寫）或 Word 副檔名
      （.doc/.docx，不分大小寫）快速判定，不需額外 OCR。
    - 4_技師執業執照：需先擷取 PDF/圖片內文字（含 OCR 備援），內文包含
      「技師執業執照」或「執業執照」才成立。
    - 5_技師會員證：內文「同時」包含「會員證」與「技師公會」才成立
      （例如：台北市水利技師公會、台灣省水利技師公會），避免景觀學會等
      一般協會會員證被誤判。
    - 6_投保證明 / 2_學歷：內容結構化程度低，維持以檔名關鍵字判斷。
    - 其餘：歸為 3_證照，不進行過期告警。

    extracted_text 僅在實際執行過 PDF/OCR 擷取時回傳（CV／未擷取則為 None），
    供效期檢核複用，避免對同一檔案重複 OCR。
    """
    fname = os.path.basename(filepath)
    fname_lower = fname.lower()
    ext = Path(filepath).suffix.lower()

    if any(kw.lower() in fname_lower for kw in CV_FILENAME_KEYWORDS):
        return "1_CV", None
    if ext in (".doc", ".docx"):
        return "1_CV", None

    pdf_path = filepath if ext == ".pdf" else convert_to_pdf(filepath, workdir)
    text = extract_pdf_text(pdf_path) if pdf_path else ""

    if "技師執業執照" in text or "執業執照" in text or "执业执照" in text:
        return "4_技師執業執照", text

    has_member_kw = ("會員證" in text) or ("会员证" in text)
    has_guild_kw = ("技師公會" in text) or ("技师公会" in text)
    if has_member_kw and has_guild_kw:
        return "5_技師會員證", text

    hint = classify_category_by_filename_hint(fname_lower)
    if hint:
        return hint, text

    return "3_證照", text


# ============================================================
# 證照效期檢核
# ============================================================

DATE_RANGE_PATTERN = re.compile(
    r"(?:自)?民國\s*(\d{2,3})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日\s*"
    r"(?:止|至)\s*(?:民國\s*)?(\d{2,3})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日"
)

# 「115年會員證」／「115年度會員證」等只標示單一年度、無明確起訖日的格式：
# 推算至該年12/31止（僅供 5_技師會員證 使用）
YEAR_ONLY_PATTERN = re.compile(r"(\d{2,3})\s*年度?\s*(?:會員證|会员证)")

# 備援：專門抓取「至／~／～」或緊接在其後、直到「止」為止的截止日期。
# 用於 OCR 誤判、格式跑版等導致「起訖日全段」比對失敗，但截止日文字仍可辨識的情況
# （如「羅翊軒」「潘冠愷」執照因掃描雜訊導致起始日段落抓不到）。
END_DATE_ONLY_PATTERN = re.compile(
    r"(?:至|~|～)\s*(?:民國\s*)?(\d{2,3})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日\s*(?:止)?"
)


def parse_license_expiry(category: str, text: str) -> dict:
    """依「已擷取好」的內文文字解析證照效期，僅供 4_技師執業執照 與
    5_技師會員證 呼叫；不重複進行 PDF/OCR 擷取。日期一律以呼叫當下的
    date.today() 動態計算，避免因常數在應用程式啟動時凍結而產生誤判。

    優先順序：明確起訖日期 > 單一年度（僅會員證，推算至當年12/31）
    > 僅截止日備援（聚焦「至/~/止」後方日期，因應 OCR 起始日段落判讀失敗）。
    """
    if not text or not text.strip():
        return {"狀態": "⚠️ 無法自動判讀（掃描檔/OCR不可用）", "起始日": "", "截止日": "",
                "備註": "請人工確認"}

    start = end = None

    m = DATE_RANGE_PATTERN.search(text)
    if m:
        ry1, rm1, rd1, ry2, rm2, rd2 = map(int, m.groups())
        try:
            start = date(ry1 + 1911, rm1, rd1)
            end = date(ry2 + 1911, rm2, rd2)
        except ValueError:
            return {"狀態": "⚠️ 日期格式錯誤", "起始日": "", "截止日": "", "備註": "請人工確認"}
    elif category == "5_技師會員證":
        m2 = YEAR_ONLY_PATTERN.search(text)
        if m2:
            ry = int(m2.group(1))
            try:
                start = date(ry + 1911, 1, 1)
                end = date(ry + 1911, 12, 31)
            except ValueError:
                return {"狀態": "⚠️ 日期格式錯誤", "起始日": "", "截止日": "", "備註": "請人工確認"}

    if end is None:
        m3 = END_DATE_ONLY_PATTERN.search(text)
        if m3:
            ry2, rm2, rd2 = map(int, m3.groups())
            try:
                end = date(ry2 + 1911, rm2, rd2)
                start = None  # 起始日因 OCR 判讀失敗而未知，僅顯示截止日
            except ValueError:
                return {"狀態": "⚠️ 日期格式錯誤", "起始日": "", "截止日": "", "備註": "請人工確認"}

    if end is None:
        return {"狀態": "⚠️ 未偵測到效期文字", "起始日": "", "截止日": "", "備註": "請人工確認"}

    today = date.today()
    days_left = (end - today).days
    if days_left < 0:
        status = "🔴 已過期"
        remark = f"逾期 {-days_left} 天"
    elif days_left <= 90:
        status = "🟡 即將過期（90天內）"
        remark = f"剩 {days_left} 天"
    else:
        status = "🟢 有效"
        remark = f"剩 {days_left} 天"

    return {"狀態": status, "起始日": start.isoformat() if start else "（OCR未判讀）",
            "截止日": end.isoformat(), "備註": remark}


# ============================================================
# 投保年資精準離線擷取（完全不消耗 AI 額度）
# ============================================================

INSURANCE_YEARS_PATTERN = re.compile(r"勞保投保年資\s*[:：]\s*(\d+)\s*年\s*(\d+)\s*日")


def parse_insurance_years(text: str):
    """從投保證明內文（如「勞保投保年資：6年 309日（截至115/07/22止）」）
    離線解析年資，完全不呼叫任何 AI API。

    換算規則：日數 >= 180 則年數 +1（四捨五入進位）；否則直接取年數。
    找不到符合格式的文字則回傳 None，由呼叫端改用 Excel 資歷欄位或預設為 0。
    """
    if not text:
        return None
    m = INSURANCE_YEARS_PATTERN.search(text)
    if not m:
        return None
    years = int(m.group(1))
    days = int(m.group(2))
    if days >= 180:
        years += 1
    return years


# ============================================================
# Badge / Layer 判定
# ============================================================

# ============================================================
# Badge / Layer 判定
# ============================================================

def clean_license_text(raw: str) -> str:
    """證照/資格候選文字專用清洗：比 clean_expertise_text 寬鬆，保留常見的
    「()（）」「-」等符號（如「國際專案管理師(PMP)」），只清除項目符號、
    控制字元等亂碼。"""
    if not raw:
        return ""
    text = BULLET_NOISE_PATTERN.sub(" ", raw)
    text = NON_TEXT_NOISE_PATTERN.sub(" ", text)
    text = re.sub(r"[^\u4e00-\u9fffA-Za-z0-9()（）\-、,，/\s]", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def _build_spaced_header_pattern(keywords: list) -> "re.Pattern":
    """將關鍵字組成「錨定於行首」的表頭正則，並容許中文字之間夾雜半形/全形
    空白（如「資 格：」），依關鍵字長度由長到短比對，避免「資格」搶先命中
    「資格及訓練」的一部分。"""
    parts = []
    for kw in keywords:
        spaced = r"[\s\u3000]*".join(re.escape(ch) for ch in kw)
        parts.append((len(kw), spaced))
    parts.sort(key=lambda x: x[0], reverse=True)
    pattern_str = "|".join(p for _, p in parts)
    return re.compile(r"^[\s\u3000]*(?:" + pattern_str + r")[\s\u3000]*[:：]?[\s\u3000]*")


# 「資格區塊」起訖邊界關鍵字：僅擷取起始標頭到下一個段落標頭之間的內容，
# 嚴禁掃描「工作經歷」等敘述性段落，避免將「擔任設計技師」「監造技師」等
# 專案說明文字誤判為證照。
QUALIFICATION_START_KEYWORDS = ["資格及訓練", "專業資格", "資格", "證照"]
QUALIFICATION_END_KEYWORDS = ["教育訓練", "工作經歷", "專案經歷", "經歷"]
QUALIFICATION_START_PATTERN = _build_spaced_header_pattern(QUALIFICATION_START_KEYWORDS)
QUALIFICATION_END_PATTERN = _build_spaced_header_pattern(QUALIFICATION_END_KEYWORDS)

LIST_MARKER_PATTERN = re.compile(r"^(?:[0-9０-９]+[.、)]|[-－•▪◆■●○])\s*")
_LICENSE_HEADER_KEYWORDS = ["學歷", "學位", "證照", "证照", "技師", "執照", "檢定",
                            "考試及格", "資格", "專業資格", "專業認證", "教育訓練",
                            "訓練合格", "會員證", "技術士"]
LICENSE_HEADER_PATTERN = re.compile(
    r"^(?:" + "|".join(sorted(_LICENSE_HEADER_KEYWORDS, key=len, reverse=True)) + r")"
    r"[\s\u3000]*[:：]?[\s\u3000]*"
)


def extract_qualification_section_text(cv_text: str) -> str:
    """區塊邊界精準擷取（Section-based Extraction）：找到「資格：」「資 格：」
    「專業資格：」「證照：」「資格及訓練：」等起始標頭，一路擷取到下一個
    「教育訓練：」「工作經歷：」「經歷：」「專案經歷：」等段落標頭為止，
    只回傳這個區塊「內部」的文字。找不到明確起始標頭則回傳空字串，交由
    呼叫端退回較寬鬆的關鍵字比對邏輯。"""
    if not cv_text:
        return ""
    lines = cv_text.splitlines()

    start_idx = None
    start_remainder = ""
    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            continue
        m = QUALIFICATION_START_PATTERN.match(stripped)
        if m:
            start_idx = i
            start_remainder = stripped[m.end():].strip()
            break
    if start_idx is None:
        return ""

    collected = [start_remainder] if start_remainder else []
    for j in range(start_idx + 1, len(lines)):
        stripped = lines[j].strip()
        if not stripped:
            continue
        if QUALIFICATION_END_PATTERN.match(stripped):
            break
        collected.append(stripped)
        if len(collected) >= 60:  # 安全上限：避免區塊未正確結束時無限往下吃
            break

    return "\n".join(collected)


def extract_badge_relevant_text(cv_text: str) -> str:
    """僅擷取 CV 中與「學歷／證照／技師／資格」相關的段落，限縮候選範圍，
    避免掃描到專案名稱等無關內容（例如專案名稱含「景觀」但當事人並非景觀技師）。
    做為 extract_qualification_section_text() 找不到明確區塊標頭時的退而
    求其次備援。"""
    if not cv_text:
        return ""
    section_keywords = ["學歷", "學位", "證照", "证照", "技師", "執照", "檢定",
                        "考試及格", "資格", "專業資格", "專業認證", "教育訓練",
                        "訓練合格", "會員證", "技術士"]
    relevant = []
    for line in cv_text.splitlines():
        line_stripped = line.strip()
        if not line_stripped:
            continue
        if any(kw in line_stripped for kw in section_keywords):
            relevant.append(line_stripped)
    return "\n".join(relevant)


def extract_license_candidates_offline(cv_text: str, cats: dict, name: str = "") -> list:
    """100% 離線：從 CV「資格」區塊，擷取同仁擁有的所有「證照/資格全稱」候選
    清單。**不做任何自動縮寫對應**——是否採用（勾選）與對應縮寫代碼，完全
    交由使用者於 Streamlit UI 手動指定。

    來源（優先序）：
    1. 「區塊邊界精準擷取」：僅擷取「資格：」～下一個「教育訓練/工作經歷/
       經歷/專案經歷」標頭之間的內容，嚴禁掃描工作經歷等敘述性段落。
    2. 找不到明確資格區塊標頭時，退回較寬鬆的關鍵字行比對
       （extract_badge_relevant_text）。
    3. 已分類為「4_技師執業執照」／「5_技師會員證」的證照檔名，各自視為
       一筆候選（檔名去除副檔名、去除人員姓名後作為顯示名稱）。
    回傳去重後的候選字串清單，保留原始出現順序。
    """
    candidates = []
    seen = set()

    def _add(text: str):
        text = (text or "").strip(" -－、,，;；/")
        if not text or len(text) < 2 or len(text) > 40:
            return
        if text in seen:
            return
        seen.add(text)
        candidates.append(text)

    section_text = extract_qualification_section_text(cv_text)
    if not section_text.strip():
        section_text = extract_badge_relevant_text(cv_text)

    for line in section_text.splitlines():
        line = LIST_MARKER_PATTERN.sub("", line.strip())
        cleaned = LICENSE_HEADER_PATTERN.sub("", line).strip()
        if not cleaned:
            continue
        for part in re.split(r"[、,，;；/]+", cleaned):
            _add(clean_license_text(part))

    for cat_key in ("4_技師執業執照", "5_技師會員證"):
        for fp in cats.get(cat_key, []):
            stem = Path(fp).stem
            if name:
                stem = stem.replace(name, "")
            stem = clean_license_text(stem).strip(" _-（）()")
            _add(stem)

    return candidates



def determine_layer(role_text: str) -> str:
    """判斷 Layer，相容「畫／劃」異體字（如：計畫主持人／計劃主持人）。
    比對順序刻意將較具體的關鍵字（如「副主持人」）放在通用「主持人」之前，
    避免被較寬鬆的規則搶先命中。"""
    role_text = str(role_text or "")

    if any(kw in role_text for kw in ["設計負責人", "副計畫主持人", "副計劃主持人"]):
        return "SubTop"
    if any(kw in role_text for kw in ["計畫主持人", "計劃主持人", "協同主持人", "代表廠商", "主持人"]):
        return "Top"
    if any(kw in role_text for kw in ["計畫顧問", "計劃顧問", "品質督導", "顧問"]):
        return "Advisor"
    if any(kw in role_text for kw in ["計畫經理", "計劃經理", "專案經理"]):
        return "Middle"
    if any(kw in role_text for kw in ["組長", "隊長"]):
        return "GroupLeader"
    if any(kw in role_text for kw in ["協力廠商", "分包"]):
        return "Subcontractor"
    return "GroupMember"


# ============================================================
# Phase 1：100% 離線欄位預填（不消耗 API，零失敗保底）
# ============================================================

DEGREE_PRIORITY_KEYWORDS = ["博士", "碩士", "研究所", "學士", "大學"]

COMPANY_HINT_KEYWORDS = ["股份有限公司", "有限公司", "工程顧問"]
TITLE_HINT_KEYWORDS = ["職稱", "職務"]

# 公司名稱結尾關鍵字（依常見長度排序，用於切分「公司全名」與「職稱」）
COMPANY_SUFFIX_KEYWORDS = [
    "股份有限公司", "有限公司", "技師事務所", "建築師事務所", "事務所",
    "分公司", "公司", "協會", "學會", "基金會",
]

# 現職／服務單位表頭：可相容「現    職：」「現　職：」等夾帶多個半形/全形空白、
# Tab 的情況（\s 已涵蓋半形空白與 Tab，額外加入 \u3000 涵蓋全形空白）
POSITION_HEADER_PATTERN = re.compile(
    r"(?:現[\s\u3000]*職|現[\s\u3000]*任|服務[\s\u3000]*單位|服務[\s\u3000]*機構|"
    r"任職[\s\u3000]*公司)[\s\u3000]*[:：\-－]?[\s\u3000]*"
)
# 表頭關鍵字黑名單：這些字樣本身不能被當作 Title/Company 的值
HEADER_BLACKLIST = {"職稱", "服務單位", "現職", "職務", "公司", "服務機構"}
# 換段落關鍵字：多行掃描現職內容時，遇到這些開頭即停止併行，避免吃進下一個段落
SECTION_BREAK_PATTERN = re.compile(r"^(?:學[\s\u3000]*歷|證照|經歷|專長|語言能力)[\s\u3000]*[:：]")

# 常用職稱字典：公司/職稱切分未能命中時，於 CV 前 15 行掃描補全
COMMON_TITLE_DICTIONARY = [
    "正工程師", "水利工程師", "監造主任", "資深協理", "協理", "副理",
    "經理", "工程師", "技師", "主任", "副總經理", "總經理", "組長", "課長", "襄理",
]

JOB_DESCRIPTION_TEMPLATES = [
    (["計畫主持人", "計劃主持人", "主持人"], "負責本專案整體規劃、品質控管、進度督導與跨單位溝通協調"),
    (["協同主持人"], "協助計畫主持人統籌專案執行，負責跨組別協調與技術審查"),
    (["計畫顧問", "計劃顧問", "顧問"], "提供專業技術諮詢與審查意見，協助提升專案品質與可行性"),
    (["品質督導"], "負責專案品質管理制度督導與稽核，確保符合契約與法規要求"),
    (["設計負責人"], "負責工程設計規劃、技術審查與設計圖說品質把關"),
    (["計畫經理", "計劃經理", "專案經理"], "負責專案進度管控、資源調度、預算執行與跨部門協調"),
    (["組長", "隊長"], "負責所屬工作組別之任務分派、技術執行與進度回報"),
    (["協力廠商", "分包"], "配合本案提供專業技術服務與資源支援"),
]

# 組別關鍵字 -> 該組別典型工作內容片語，用於讓 JobDescription 依 GroupName 動態變化
GROUP_NAME_HINTS = [
    (["現場調查", "現勘"], "辦理現場勘查與基礎資料蒐集"),
    (["細部設計", "設計"], "執行細部設計與技術規範撰擬"),
    (["監造", "監督施工"], "辦理工程監造與施工品質查核"),
    (["測量"], "執行測量作業與圖籍檢核"),
    (["環境", "生態"], "辦理環境影響評估與生態調查"),
    (["水理", "水文"], "執行水理水文分析與模擬"),
    (["行政", "文書"], "辦理專案行政聯繫與文書作業"),
]


def clean_degree_text(text) -> str:
    """清除 Degree 欄位常見的「學歷：」／「學 歷：」等字頭。可正確處理全形
    空白（\\u3000）、半形空白，以及全形「：」與半形「:」冒號的各種組合
    （如 CV/OCR 文字常見的「學 歷：國立臺灣大學…」）。"""
    text = str(text or "")
    text = re.sub(r"^\s*學[\s\u3000]*歷[\s\u3000]*[:：]\s*", "", text)
    return text.strip()


def extract_degree_offline(cv_text: str) -> str:
    """離線正則：依優先序（博士＞碩士＞研究所＞學士＞大學）搜尋 CV 中含學歷
    關鍵字的段落，並強制清除「學歷：」等字頭，只保留校名／科系／學位本身。"""
    if not cv_text:
        return ""
    lines = [l.strip() for l in cv_text.splitlines() if l.strip()]
    for kw in DEGREE_PRIORITY_KEYWORDS:
        for line in lines:
            if kw in line:
                return clean_degree_text(line)[:60]
    return ""


def split_company_title(text: str):
    """依公司名稱結尾關鍵字（如：分公司／股份有限公司／事務所）切分「公司全名」
    與「職稱」，取結尾關鍵字中「結束位置最晚、且同結束位置下最長」者，避免公司
    名稱中途出現的「工程顧問」等詞被誤判為結尾。"""
    text = text.strip()
    if not text:
        return "", ""
    window = text[:40]
    best_end, best_kw = -1, None
    for kw in COMPANY_SUFFIX_KEYWORDS:
        idx = window.find(kw)
        if idx == -1:
            continue
        end = idx + len(kw)
        if end > best_end or (end == best_end and (best_kw is None or len(kw) > len(best_kw))):
            best_end, best_kw = end, kw

    if best_end != -1:
        company = text[:best_end].strip()
        title = text[best_end:].strip()
        title = title.strip(" ,，、-:：")
        return company, title

    # 找不到公司關鍵字：以第一個空白切分（公司名稱 職稱）
    parts = text.split(None, 1)
    if len(parts) == 2:
        return parts[0], parts[1]
    return text, ""


def extract_company_title_offline(cv_text: str):
    """離線正則：搜尋「現職／服務單位／職稱」等關鍵字段落。

    - 相容「現    職：」夾帶多個半形/全形空白或 Tab 的情況（POSITION_HEADER_PATTERN）。
    - 支援多行現職結構：表頭與內容分行時（如「現職：」單獨一行，下一行才是
      「-美商傑明工程顧問(股)台灣分公司　　工程師」），自動向下掃描最多 3 行，
      去除開頭的「-」項目符號後再精準切分公司全名與職稱。
    - 常用職稱字典備援：若切分未抓到職稱，改從 CV 前 15 行掃描常見職稱字樣
      （正工程師／水利工程師／監造主任／資深協理／工程師…）。
    - 全程禁止把「職稱」「服務單位」等表頭字樣本身填入欄位。
    """
    if not cv_text:
        return "", ""
    lines = [l.strip() for l in cv_text.splitlines()]
    company, title = "", ""

    for idx, line in enumerate(lines):
        if not line:
            continue
        m = POSITION_HEADER_PATTERN.match(line)
        if not m:
            continue

        remainder = line[m.end():].strip()
        candidates_text = [remainder] if remainder else []

        j = idx + 1
        lookahead = 0
        while j < len(lines) and lookahead < 3:
            nxt = lines[j]
            if not nxt:
                j += 1
                continue
            if SECTION_BREAK_PATTERN.match(nxt):
                break
            candidates_text.append(nxt)
            lookahead += 1
            j += 1
            if candidates_text:
                break  # 湊到第一組非空內容即可嘗試切分，避免誤吃進不相關段落

        full_text = " ".join(t for t in candidates_text if t).strip()
        full_text = re.sub(r"^[-－•]\s*", "", full_text)  # 去除開頭的項目符號
        if not full_text or full_text in HEADER_BLACKLIST:
            continue

        c, t = split_company_title(full_text)
        if c and c not in HEADER_BLACKLIST and not company:
            company = c
        if t and t not in HEADER_BLACKLIST and not title:
            title = t

        if company:
            break

    if not company:
        for line in lines:
            if any(kw in line for kw in COMPANY_HINT_KEYWORDS):
                cleaned = POSITION_HEADER_PATTERN.sub("", line).strip()
                c, _ = split_company_title(cleaned)
                if c and c not in HEADER_BLACKLIST:
                    company = c
                    break

    if not title:
        for line in lines[:15]:
            for t_kw in COMMON_TITLE_DICTIONARY:
                if t_kw in line:
                    title = t_kw
                    break
            if title:
                break

    if not title:
        for line in lines:
            if any(kw in line for kw in TITLE_HINT_KEYWORDS):
                parts = re.split(r"[:：]", line, maxsplit=1)
                if len(parts) > 1 and parts[1].strip():
                    candidate = parts[1].strip()[:20]
                    if candidate not in HEADER_BLACKLIST:
                        title = candidate
                        break

    return company, title


BULLET_NOISE_PATTERN = re.compile(r"[•▪◆■●○\uf06f\uf0b7\u2022\u25cf\*]+")
NON_TEXT_NOISE_PATTERN = re.compile(r"[\x00-\x1f\x7f-\x9f]+")


def clean_expertise_text(raw: str) -> str:
    """清除項目符號、控制字元／亂碼與非常見標點，只保留中英文數字與常見分隔符。"""
    if not raw:
        return ""
    text = BULLET_NOISE_PATTERN.sub(" ", raw)
    text = NON_TEXT_NOISE_PATTERN.sub(" ", text)
    text = re.sub(r"[^\u4e00-\u9fffA-Za-z0-9、,，/\s]", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def extract_expertise_offline(cv_text: str, license_candidates: list = None) -> str:
    """離線正則：優先從 CV「專長」段落擷取關鍵字，清除亂碼／項目符號後，
    僅保留 2~10 字的精簡詞彙（過濾長句與雜訊），統一以「/」分隔輸出 4~6 項；
    找不到則以已擷取到的證照/資格候選清單（前幾筆）展開為專長描述，確保
    欄位不為空。"""
    if cv_text:
        for kw in ["專長", "专长", "專業領域", "专业领域"]:
            idx = cv_text.find(kw)
            if idx != -1:
                segment = cv_text[idx: idx + 150]
                parts = re.split(r"[:：]", segment, maxsplit=1)
                value = parts[1] if len(parts) > 1 else parts[0]
                value = clean_expertise_text(value)
                tokens = [t.strip() for t in re.split(r"[、,，/\s]+", value) if t.strip()]
                tokens = [t for t in tokens if 2 <= len(t) <= 10]
                if len(tokens) >= 2:
                    return "/".join(tokens[:6])

    if license_candidates:
        trimmed = [c[:12] for c in license_candidates[:4]]
        if trimmed:
            return "/".join(trimmed)
    return ""


def extract_representative_projects(cv_text: str, max_n: int = 2) -> str:
    """離線正則：從 CV 中挑出 1~2 項含年份且含「專案/計畫/工程/案」字樣的
    代表性經歷，清洗雜訊後回傳精簡片語（供 BioNarrative 與 AI 草稿共用）。"""
    if not cv_text:
        return ""
    candidates = []
    for line in cv_text.splitlines():
        line = line.strip()
        if not line or len(line) > 60:
            continue
        if re.search(r"\d{2,4}\s*年", line) and any(k in line for k in ["專案", "計畫", "計劃", "工程", "案"]):
            cleaned = re.sub(r"^\d{2,4}\s*年\s*", "", line)
            cleaned = clean_expertise_text(cleaned)
            if cleaned:
                candidates.append(cleaned[:30])

    seen, uniq = set(), []
    for c in candidates:
        if c not in seen:
            seen.add(c)
            uniq.append(c)
    return "、".join(uniq[:max_n])


WORK_CONTENT_SECTION_KEYWORDS = ["工作內容", "工作內容摘要", "職務內容", "專案職責",
                                 "主要職責", "工作經歷", "負責事項", "職掌", "經歷"]
WORK_CONTENT_ACTION_VERBS = ["負責", "辦理", "執行", "規劃", "設計", "監造", "管理",
                             "協調", "審查", "建置", "分析", "督導", "撰擬", "查核"]
WORK_CONTENT_SECTION_END_KEYWORDS = ["學歷", "證照", "語言能力", "技能", "專長"]


def extract_work_content_offline(cv_text: str, max_sentences: int = 3) -> str:
    """離線正則：從 CV「經歷／工作內容／專案職責」等段落擷取 2~3 句本人
    實際負責的核心工作內容（含常見動詞如「負責/辦理/執行/規劃」等），供撰寫
    JobDescription 與 BioNarrative 的實務素材依據，並併入送交 AI 的草稿中。"""
    if not cv_text:
        return ""

    candidates = []
    capture = False
    for raw_line in cv_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        if any(kw in line for kw in WORK_CONTENT_SECTION_KEYWORDS):
            capture = True
            parts = re.split(r"[:：]", line, maxsplit=1)
            if len(parts) > 1 and parts[1].strip():
                cleaned = clean_expertise_text(parts[1].strip())
                if cleaned:
                    candidates.append(cleaned)
            continue

        if capture:
            if len(line) < 15 and any(kw in line for kw in WORK_CONTENT_SECTION_END_KEYWORDS):
                capture = False
                continue
            cleaned = clean_expertise_text(line)
            if cleaned and any(verb in cleaned for verb in WORK_CONTENT_ACTION_VERBS):
                candidates.append(cleaned[:60])

        if len(candidates) >= max_sentences:
            break

    seen, uniq = set(), []
    for c in candidates:
        if c not in seen:
            seen.add(c)
            uniq.append(c)
        if len(uniq) >= max_sentences:
            break

    return "；".join(uniq)


def build_fallback_bio(name: str, degree: str, company: str, title: str, years,
                        project_snippet: str, expertise: str) -> str:
    """BioNarrative 保底樣板：比照標案範本風格採四段式結構——
    ①學歷與現職 ②代表性經歷 ③核心專長 ④履約效益，確保即使離線也具備說服力，
    不會是單薄的單句罐頭語句。"""
    degree_part = degree or "相關專業"
    company_part = company or "本公司"
    title_part = title or "工程顧問"
    years_part = f"{years}年" if years else "多年"

    part1 = f"{name}君具{years_part}相關工程資歷，現任{company_part}{title_part}，具備{degree_part}學歷。"
    if project_snippet:
        part2 = f"曾參與{project_snippet}等代表性專案，實務經驗豐富。"
    else:
        part2 = "曾參與多項公共工程規劃、設計與監造專案，實務經驗豐富。"
    part3 = f"核心專長涵蓋{expertise}。" if expertise else "核心專長涵蓋水利與工程專案管理。"
    part4 = "具備良好團隊協調與品質控管能力，能確保本案高品質履約。"

    return part1 + part2 + part3 + part4


def build_fallback_job_description(role_text: str, group_name: str = "", work_content: str = "") -> str:
    """JobDescription 保底樣板：依團隊職務（Role）＋部門/組別（GroupName）＋
    Python 從 CV 萃取出的實際工作內容三者動態組合，避免所有組員/組長套用完全
    相同的罐頭句子。優先呈現 CV 實務內容，找不到才退回角色/組別樣板。"""
    role_text = str(role_text or "").strip()
    group_name = str(group_name or "").strip()
    if group_name in ("", "—", "nan", "None"):
        group_name = ""
    work_content = (work_content or "").strip()

    base = None
    for kws, desc in JOB_DESCRIPTION_TEMPLATES:
        if any(kw in role_text for kw in kws):
            base = desc
            break
    if base is None:
        base = "負責所屬工作項目之執行、協調與品質把關，配合專案進度完成各階段任務"

    if work_content:
        prefix = f"於{group_name}擔任{role_text or '團隊成員'}，" if group_name else \
                 (f"擔任{role_text}，" if role_text else "")
        combined = f"{prefix}{work_content}"
        return combined[:50] if len(combined) > 50 else combined

    group_hint = ""
    if group_name:
        for kws, hint in GROUP_NAME_HINTS:
            if any(kw in group_name for kw in kws):
                group_hint = hint
                break

    if group_name and group_hint:
        return f"於{group_name}擔任{role_text or '團隊成員'}，{group_hint}，{base}。"
    if group_name:
        return f"於{group_name}擔任{role_text or '團隊成員'}，{base}。"
    return base + "。"



# ============================================================
# PDF 合併
# ============================================================

def merge_person_pdfs(ordered_names: list, files_by_person: dict, workdir: str) -> bytes:
    if PdfWriter is None:
        raise RuntimeError("pypdf 套件未安裝，無法合併 PDF")

    writer = PdfWriter()
    page_cursor = 0

    for person in ordered_names:
        cats = files_by_person.get(person, {})
        person_start_page = page_cursor
        any_page = False

        for cat in CATEGORY_ORDER:
            filelist = sorted(cats.get(cat, []))
            for fp in filelist:
                pdf_path = convert_to_pdf(fp, workdir)
                if not pdf_path:
                    continue
                try:
                    reader = PdfReader(pdf_path)
                    for page in reader.pages:
                        writer.add_page(page)
                        page_cursor += 1
                        any_page = True
                except Exception:
                    continue

        if any_page:
            try:
                writer.add_outline_item(person, person_start_page)
            except Exception:
                pass

    buffer = io.BytesIO()
    writer.write(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ============================================================
# Excel 輸出
# ============================================================

def build_template_excel(df: pd.DataFrame) -> bytes:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Staffing"
    wb.calculation.fullCalcOnLoad = True

    header_font = Font(name="微軟正黑體", bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="2C3E50", end_color="2C3E50", fill_type="solid")
    body_font = Font(name="微軟正黑體")

    for col_idx, col_name in enumerate(TEMPLATE_COLUMNS, start=1):
        cell = ws.cell(row=1, column=col_idx, value=col_name)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center")

    safe_df = df.copy()
    for col in TEMPLATE_COLUMNS:
        if col not in safe_df.columns:
            safe_df[col] = ""
    safe_df = safe_df[TEMPLATE_COLUMNS]

    for row_idx, row in enumerate(safe_df.itertuples(index=False), start=2):
        for col_idx, value in enumerate(row, start=1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.font = body_font
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    widths = [10, 14, 12, 10, 14, 20, 35, 10, 16, 10, 22, 28, 28, 45]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = "A2"

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def build_license_report_excel(df_license: pd.DataFrame) -> bytes:
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df_license.to_excel(writer, index=False, sheet_name="證照效期檢核")
        wb = writer.book
        wb.calculation.fullCalcOnLoad = True
        ws = writer.sheets["證照效期檢核"]

        header_font = Font(name="微軟正黑體", bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="2C3E50", end_color="2C3E50", fill_type="solid")
        for cell in ws[1]:
            cell.font = header_font
            cell.fill = header_fill

        red_fill = PatternFill(start_color="FADBD8", end_color="FADBD8", fill_type="solid")
        yellow_fill = PatternFill(start_color="FCF3CF", end_color="FCF3CF", fill_type="solid")

        status_col_idx = None
        for idx, cell in enumerate(ws[1], start=1):
            if cell.value == "狀態":
                status_col_idx = idx
                break

        if status_col_idx and ws.max_row >= 2:
            for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
                val = row[status_col_idx - 1].value
                if val and "已過期" in str(val):
                    for c in row:
                        c.fill = red_fill
                elif val and "即將過期" in str(val):
                    for c in row:
                        c.fill = yellow_fill

        for column_cells in ws.columns:
            length = max((len(str(c.value)) if c.value else 0) for c in column_cells)
            col_letter = column_cells[0].column_letter
            ws.column_dimensions[col_letter].width = min(max(length + 2, 10), 40)

    buffer.seek(0)
    return buffer.getvalue()


# ============================================================
# 主流程（Phase 1：掃描與比對 / Phase 2：確認後執行）
# ============================================================

def scan_and_match_files(excel_file, zip_file) -> dict:
    """Phase 1：讀取人員名單、解壓 ZIP，並以三段式雙軌配對機制為每個檔案
    產生高/中/低信心度配對結果。不執行 AI 呼叫、PDF 合併等重工作，
    以便使用者能先確認/修正配對結果。"""
    df_people = pd.read_excel(excel_file)
    required_cols = ["順序", "姓名", "部門", "團隊職務"]
    missing = [c for c in required_cols if c not in df_people.columns]
    if missing:
        raise ValueError(f"人員名單 Excel 缺少必要欄位：{'、'.join(missing)}")
    if "英文姓名" not in df_people.columns:
        df_people["英文姓名"] = ""

    df_people = df_people.sort_values("順序").reset_index(drop=True)
    names = [str(n) for n in df_people["姓名"].tolist()]
    aliases = build_alias_list(df_people)

    workdir = tempfile.mkdtemp(prefix="staffing_")
    extract_dir = os.path.join(workdir, "extracted")
    convert_workdir = os.path.join(workdir, "converted")
    os.makedirs(extract_dir, exist_ok=True)
    os.makedirs(convert_workdir, exist_ok=True)

    with zipfile.ZipFile(zip_file) as z:
        z.extractall(extract_dir)

    all_files = []
    for root, _dirs, files in os.walk(extract_dir):
        root_lower = root.lower()
        if any(token.lower() in root_lower for token in EXCLUDE_DIR_TOKENS):
            continue
        for f in files:
            if f.startswith("."):
                continue
            all_files.append(os.path.join(root, f))

    high_matches = {}      # filepath -> (person, score, matched_alias_text)
    pending_matches = []   # 中信心度，待使用者確認
    low_matches = []       # 低信心度，待使用者手動指定

    for fp in all_files:
        fname = os.path.basename(fp)
        ranked = match_person_candidates(fname, aliases)
        if ranked:
            top_person, top_score, top_alias = ranked[0]
        else:
            top_person, top_score, top_alias = None, 0.0, ""

        if top_score >= HIGH_CONFIDENCE_THRESHOLD:
            high_matches[fp] = (top_person, top_score, top_alias)
        elif top_score >= MEDIUM_CONFIDENCE_THRESHOLD:
            pending_matches.append({
                "file": fp, "fname": fname,
                "candidates": ranked[:5], "default": top_person,
            })
        else:
            low_matches.append({
                "file": fp, "fname": fname, "candidates": ranked[:5],
            })

    return {
        "df_people": df_people,
        "names": names,
        "workdir": workdir,
        "extract_dir": extract_dir,
        "convert_workdir": convert_workdir,
        "high_matches": high_matches,
        "pending_matches": pending_matches,
        "low_matches": low_matches,
    }


def finalize_processing(scan: dict, pending_selection: dict, low_selection: dict) -> dict:
    """Phase 2：整合使用者確認後的人員配對結果，以「內文導向」對每個檔案
    進行分類與 OCR 掃描（含進度條），執行證照效期檢核、100% 離線欄位萃取，
    並產出合併附錄 PDF。全程不呼叫任何外部 API。"""
    df_people = scan["df_people"]
    names = scan["names"]
    convert_workdir = scan["convert_workdir"]

    final_mapping = {fp: info[0] for fp, info in scan["high_matches"].items()}
    for item in scan["pending_matches"]:
        person = pending_selection.get(item["file"])
        if person:
            final_mapping[item["file"]] = person

    still_unmatched = []
    for item in scan["low_matches"]:
        person = low_selection.get(item["file"])
        if person:
            final_mapping[item["file"]] = person
        else:
            still_unmatched.append({"file": item["file"], "fname": item["fname"]})

    # ------------------------------------------------------------
    # Phase 2a：內文導向分類 + OCR 掃描（含動態進度條）
    # ------------------------------------------------------------
    files_by_person = {}
    license_rows = []
    insurance_years_by_person = {}  # 姓名 -> 離線解析出的投保年資（優先權高於 AI）

    file_items = list(final_mapping.items())
    total_files = max(len(file_items), 1)
    ocr_status = st.status("正在進行內文與 OCR 掃描...", expanded=True)
    ocr_progress = st.progress(0.0)

    for idx, (fp, person) in enumerate(file_items, start=1):
        fname = os.path.basename(fp)
        ocr_status.write(f"正在進行內文與 OCR 掃描：{fname} ({idx}/{total_files})")
        ocr_progress.progress(idx / total_files)

        category, text = classify_file_content(fp, convert_workdir)
        files_by_person.setdefault(person, {}).setdefault(category, []).append(fp)

        # 逐頁掃描證照/會員證：正確處理「多張證照合併於同一份PDF」的情況
        # （例如潘冠愷合併.pdf 同時含執業執照頁與會員證頁），不會因整份檔案
        # 只被歸類為單一種類而漏掉其他頁面的證照效期。
        if category != "1_CV":
            page_results = find_license_pages(fp, convert_workdir)
            for res in page_results:
                license_rows.append({
                    "姓名": person,
                    "文件類別": "技師執業執照" if res["category"] == "4_技師執業執照" else "技師公會會員證",
                    "檔名": fname,
                    **res["check"],
                })

        if category == "6_投保證明":
            # 完全離線解析，不消耗任何 AI 額度；同一人若有多筆投保證明，取年資較大者
            offline_years = parse_insurance_years(text or "")
            if offline_years is not None:
                if person not in insurance_years_by_person or offline_years > insurance_years_by_person[person]:
                    insurance_years_by_person[person] = offline_years

    ocr_progress.progress(1.0)
    ocr_progress.empty()
    ocr_status.update(label=f"內文與 OCR 掃描完成（共 {total_files} 個檔案）", state="complete")

    # ------------------------------------------------------------
    # Phase 2b：100% 離線欄位萃取（零 API 消耗、零失敗保底）
    # ------------------------------------------------------------

    # Excel 資歷欄位（YearsOfExp 第二層備援，找不到投保證明時使用）
    excel_years_col = None
    for col in ["資歷", "年資", "YearsOfExp", "工作年資"]:
        if col in df_people.columns:
            excel_years_col = col
            break

    staffing_rows = []
    diagnostics = {}
    no_cv_people = []
    license_candidates_by_person = {}  # 姓名 -> [證照/資格全稱候選, ...]

    progress = st.progress(0.0, text="正在進行離線欄位萃取...")
    total = max(len(df_people), 1)

    for i, prow in df_people.iterrows():
        name = str(prow["姓名"])
        role = str(prow.get("團隊職務", "") or "")
        group = str(prow.get("部門", "") or "")
        progress.progress(i / total, text=f"離線欄位萃取：{name}")

        cats = files_by_person.get(name, {})
        files_by_category_named = {
            cat: [os.path.basename(f) for f in files] for cat, files in cats.items()
        }

        # CV 文字擷取：PDF（含大寫 .PDF）與 Word（含大寫 .DOC/.DOCX）皆支援，
        # 讀取「完整」全文，不做任何頁數/字數截斷
        cv_files = cats.get("1_CV", [])
        cv_text = extract_cv_text(cv_files[0], convert_workdir) if cv_files else ""
        if not cv_files:
            no_cv_people.append(name)

        # ---------- Phase 1：100% 離線預填（零 API 消耗，零失敗）----------
        # 證照/資格：僅擷取全稱候選清單，不做任何自動縮寫對應——是否採用
        # 與縮寫代碼完全交由使用者於 UI 手動指定（見 tab2 的勾選面板）。
        license_candidates = extract_license_candidates_offline(cv_text, cats, name)
        license_candidates_by_person[name] = license_candidates

        degree_offline = clean_degree_text(extract_degree_offline(cv_text))
        company_offline, title_offline = extract_company_title_offline(cv_text)
        expertise_offline = extract_expertise_offline(cv_text, license_candidates)
        project_snippet = extract_representative_projects(cv_text, max_n=2)
        work_content_offline = extract_work_content_offline(cv_text, max_sentences=3)

        # YearsOfExp：優先投保證明離線解析 > Excel 資歷欄位 > 0
        years_final = insurance_years_by_person.get(name)
        if years_final is None and excel_years_col is not None:
            raw_val = prow.get(excel_years_col)
            if pd.notna(raw_val):
                try:
                    years_final = int(float(raw_val))
                except (ValueError, TypeError):
                    years_final = None
        if years_final is None:
            years_final = 0

        # BioNarrative／JobDescription：以離線草稿完整生成（四段式／依角色+組別+
        # CV實際工作內容動態組合），內容已盡量豐富完整，可直接作為最終輸出，
        # 亦方便使用者後續自行複製到 Gemini 網頁版等工具做最終文字潤飾。
        bio_final = build_fallback_bio(name, degree_offline, company_offline, title_offline,
                                       years_final, project_snippet, expertise_offline)
        job_final = build_fallback_job_description(role, group, work_content_offline)

        final_data = {
            "Title": title_offline,
            "Company": company_offline,
            "Degree": degree_offline,
            "Expertise": expertise_offline,
            "JobDescription": job_final,
            "BioNarrative": bio_final,
        }

        if not cv_files:
            diag_status = "🔴 缺少 CV 檔案"
            diag_reason = "ZIP 內未比對到任何 CV 檔案，已套用通用樣板保底"
        elif not cv_text.strip():
            diag_status = "🟡 資料待補"
            diag_reason = "CV 檔案無法擷取到文字內容（可能為純掃描影像，OCR 失敗），已套用樣板保底"
            st.warning(f"⚠️ 「{name}」的 CV 檔案無法擷取到文字內容（可能為純掃描影像），"
                       "已套用樣板保底，可視需要改由手動填寫或確認 OCR 環境。")
        else:
            diag_status = "🟢 離線萃取完整"
            diag_reason = ""

        diagnostics[name] = {
            "status": diag_status,
            "reason": diag_reason,
            "files_by_category": files_by_category_named,
        }

        layer = determine_layer(role)

        # 初始 Licenses／Badges：預設所有候選皆勾選、縮寫皆留空
        # （精確值會在 tab2 的證照勾選面板依使用者操作即時重新計算）
        licenses_final = "、".join(license_candidates)
        badges_final = ""

        staffing_rows.append({
            "Layer": layer,
            "Role": role,
            "GroupName": group if group and group.lower() != "nan" else "—",
            "Name": name,
            "Title": final_data["Title"],
            "Company": final_data["Company"],
            "Licenses": licenses_final,
            "Badges": badges_final,
            "PhotoName": f"{name}.jpg",
            "YearsOfExp": years_final,
            "Degree": final_data["Degree"],
            "JobDescription": final_data["JobDescription"],
            "Expertise": final_data["Expertise"],
            "BioNarrative": final_data["BioNarrative"],
        })

    progress.progress(1.0, text="離線欄位萃取完成")
    progress.empty()

    if no_cv_people:
        st.warning("⚠️ 以下人員找不到 CV 檔案，已套用通用樣板保底，建議手動補充精確資訊：" +
                   "、".join(no_cv_people))

    df_staffing = pd.DataFrame(staffing_rows, columns=TEMPLATE_COLUMNS)
    df_license = pd.DataFrame(
        license_rows,
        columns=["姓名", "文件類別", "檔名", "狀態", "起始日", "截止日", "備註"],
    )

    merged_pdf_bytes = b""
    try:
        merged_pdf_bytes = merge_person_pdfs(names, files_by_person, convert_workdir)
    except Exception as e:
        st.error(f"附錄 PDF 合併發生錯誤：{e}")

    return {
        "df_staffing": df_staffing,
        "df_license": df_license,
        "merged_pdf": merged_pdf_bytes,
        "unmatched_files": still_unmatched,
        "diagnostics": diagnostics,
        "license_candidates_by_person": license_candidates_by_person,
    }


# ============================================================
# Tab2 UI 輔助函式：Mapping 診斷呈現
# ============================================================

def format_mapping_line(name: str, diag: dict) -> str:
    """組成單一人員的檔案配對摘要行，CV 一律顯示（無則標示「無」），
    其餘分類僅在有檔案時列出。"""
    files_by_cat = diag.get("files_by_category", {})
    cv_files = files_by_cat.get("1_CV", [])
    parts = ["[CV] " + ("、".join(cv_files) if cv_files else "無")]
    for cat in ["2_學歷", "3_證照", "4_技師執業執照", "5_技師會員證", "6_投保證明"]:
        files = files_by_cat.get(cat, [])
        if files:
            parts.append(f"[{CATEGORY_LABELS[cat]}] " + "、".join(files))
    return f"**{name}**：" + "｜".join(parts) + f"（狀態：{diag.get('status', '')}）"


def build_default_license_state(license_candidates_by_person: dict) -> dict:
    """依離線擷取到的證照/資格候選清單，建立預設「全部勾選」的 UI 狀態結構，
    供 tab2 的個人勾選面板讀取與編輯。縮寫代碼已改為全局 Summary 集中管理
    （見 license_abbrev_map），每個項目不再各自保存縮寫。"""
    state = {}
    for name, candidates in license_candidates_by_person.items():
        state[name] = [{"text": c, "checked": True} for c in candidates]
    return state


def get_unique_checked_licenses(license_state: dict, person_order: list = None) -> list:
    """彙整「目前所有同仁已勾選採用」的證照/資格全稱，去重複後回傳，
    依人員 Excel 順序、同一人內出現順序，保留穩定的首次出現順序。"""
    order = person_order if person_order else list(license_state.keys())
    seen = set()
    unique_texts = []
    for name in order:
        for item in license_state.get(name, []):
            if item.get("checked") and item["text"] not in seen:
                seen.add(item["text"])
                unique_texts.append(item["text"])
    return unique_texts


def compute_abbrev_stats(license_state: dict, abbrev_map: dict) -> dict:
    """統計每個縮寫代碼目前被「已勾選」證照使用的總次數（人/張數），
    用於 Summary 看板旁的即時統計（如：【技】4人｜【品】3人｜【P】2人）。"""
    license_checked_count = {}
    for items in license_state.values():
        for item in items:
            if item.get("checked"):
                license_checked_count[item["text"]] = license_checked_count.get(item["text"], 0) + 1

    stats = {}
    for text, count in license_checked_count.items():
        code = (abbrev_map.get(text) or "").strip()
        if code:
            stats[code] = stats.get(code, 0) + count
    return stats


def recompute_licenses_badges(df_staffing: pd.DataFrame, license_state: dict,
                               abbrev_map: dict) -> pd.DataFrame:
    """依目前 UI 上的勾選狀態與全局縮寫對照表，重新計算 Licenses／Badges
    兩欄並寫回 DataFrame，其餘欄位保持不變。

    - Licenses：輸出該同仁「已勾選」的證照/資格全稱，頓號分隔。
    - Badges：依 Summary 看板上「該證照全稱對應的縮寫」，僅輸出「已勾選且
      對應縮寫不為空」的代碼，半形逗號分隔（同一人重複代碼只列一次）。

    僅更新「存在於 license_state 中的人員」列，避免覆蓋使用者手動於表格中
    新增（如協力廠商 Subcontractor）之列的既有內容。"""
    if df_staffing is None or df_staffing.empty:
        return df_staffing
    df = df_staffing.copy()
    for idx, row in df.iterrows():
        name = row.get("Name")
        if name not in license_state:
            continue
        items = license_state[name]
        checked_items = [it for it in items if it.get("checked")]
        df.at[idx, "Licenses"] = "、".join(it["text"] for it in checked_items)

        badge_codes = []
        for it in checked_items:
            code = (abbrev_map.get(it["text"]) or "").strip()
            if code and code not in badge_codes:
                badge_codes.append(code)
        df.at[idx, "Badges"] = ",".join(badge_codes)
    return df


def run_finalize_and_store(scan: dict) -> None:
    """執行 Phase 2 並將結果寫回 st.session_state，供 Step 2 按鈕與
    tab2 的「套用修正並重新處理」按鈕共用。"""
    with st.spinner("處理中，請稍候（含 PDF 轉檔、OCR 效期檢核與離線欄位萃取，可能需數分鐘）..."):
        try:
            result = finalize_processing(
                scan, st.session_state.pending_selection, st.session_state.low_selection,
            )
            st.session_state.df_staffing = result["df_staffing"]
            st.session_state.df_license = result["df_license"]
            st.session_state.merged_pdf_bytes = result["merged_pdf"]
            st.session_state.unmatched_files = result["unmatched_files"]
            st.session_state.diagnostics = result["diagnostics"]
            st.session_state.license_state = build_default_license_state(
                result["license_candidates_by_person"]
            )
            # license_abbrev_map（縮寫代碼，依證照全稱為 key）為全局狀態，
            # 重新處理時保留使用者已填寫的對照，不清空。
            st.success("✅ 處理完成！請於下方分頁查看結果並下載檔案。")
        except Exception as e:
            st.error(f"❌ 處理失敗：{e}")


# ============================================================
# Streamlit 介面
# ============================================================

st.set_page_config(page_title="備標人員附錄整理、證照檢核與 CV 結構化轉檔系統", layout="wide")

st.title("📋 備標人員附錄整理、證照檢核與 CV 結構化轉檔系統")
st.caption("上傳人員名單 Excel 與證明文件 ZIP，自動排序合併附錄 PDF、檢核證照效期，"
           "並 100% 本地離線產出可直接餵給組織圖產生器的 Template_Staffing.xlsx"
           "（不呼叫任何外部 AI／雲端 API，資料不外流）")

for key, default in [
    ("scan", None),
    ("pending_selection", {}),
    ("low_selection", {}),
    ("df_staffing", None),
    ("df_license", None),
    ("merged_pdf_bytes", None),
    ("unmatched_files", []),
    ("diagnostics", {}),
    ("license_state", {}),
    ("license_abbrev_map", {}),
]:
    if key not in st.session_state:
        st.session_state[key] = default

with st.sidebar:
    st.header("① 檔案上傳")
    excel_file = st.file_uploader(
        "人員名單 Excel（需含：順序／姓名／部門／團隊職務，建議另含「英文姓名」欄）",
        type=["xlsx"],
    )
    zip_file = st.file_uploader("人員證明文件 ZIP（CV / 學歷 / 證照 / 投保等 PDF・Word）", type=["zip"])

    st.caption("🔒 本工具 100% 本地離線運作，不會呼叫任何外部 AI／雲端 API，"
               "所有人員資料與標案內容僅於本次工作階段處理，不會外流。")

    st.divider()
    scan_btn = st.button("① 掃描並比對人員檔案", use_container_width=True)

# ------------------------------------------------------------
# Step 1：掃描並比對（三段式雙軌配對機制）
# ------------------------------------------------------------

if scan_btn:
    if not excel_file or not zip_file:
        st.error("請先於左側上傳人員名單 Excel 與證明文件 ZIP。")
    else:
        with st.spinner("掃描檔案並比對人員中，請稍候..."):
            try:
                scan = scan_and_match_files(excel_file, zip_file)
                st.session_state.scan = scan
                st.session_state.pending_selection = {
                    item["file"]: item["default"] for item in scan["pending_matches"]
                }
                st.session_state.low_selection = {
                    item["file"]: None for item in scan["low_matches"]
                }
                # 重置尚未產生的最終結果
                st.session_state.df_staffing = None
                st.session_state.df_license = None
                st.session_state.merged_pdf_bytes = None
                st.session_state.unmatched_files = []
                st.session_state.diagnostics = {}
                st.session_state.license_state = {}
                st.session_state.license_abbrev_map = {}
                st.success(
                    f"✅ 掃描完成：高信心度自動配對 {len(scan['high_matches'])} 個檔案、"
                    f"待確認 {len(scan['pending_matches'])} 個、"
                    f"低信心度待手動指定 {len(scan['low_matches'])} 個。"
                )
            except Exception as e:
                st.error(f"❌ 掃描失敗：{e}")

scan = st.session_state.scan

if scan is not None:
    st.subheader("Step 1｜檔案配對結果")

    with st.expander(f"✅ 已自動配對（高信心度 ≥ {HIGH_CONFIDENCE_THRESHOLD:.0%}）："
                      f"{len(scan['high_matches'])} 個檔案", expanded=False):
        if scan["high_matches"]:
            high_rows = [
                {
                    "檔案": os.path.basename(fp),
                    "配對依據": alias_text,
                    "配對人員": person,
                    "信心度": f"{score:.0%}",
                }
                for fp, (person, score, alias_text) in scan["high_matches"].items()
            ]
            st.dataframe(pd.DataFrame(high_rows), use_container_width=True, hide_index=True)
            for row in high_rows:
                st.caption(f"已自動配對：{row['配對依據']} → {row['配對人員']}")
        else:
            st.write("（無）")

    if scan["pending_matches"]:
        st.markdown(f"### ⚠️ 待確認配對（中信心度 {MEDIUM_CONFIDENCE_THRESHOLD:.0%}～"
                    f"{HIGH_CONFIDENCE_THRESHOLD:.0%}）：{len(scan['pending_matches'])} 個檔案")
        st.caption("已預設選中最高機率的候選人員，請確認或修正後再進行 Step 2。")
        for item in scan["pending_matches"]:
            option_names = [c[0] for c in item["candidates"]]
            for n in scan["names"]:
                if n not in option_names:
                    option_names.append(n)
            options = ["— 略過此檔案 —"] + option_names
            default_person = st.session_state.pending_selection.get(item["file"], item["default"])
            default_idx = options.index(default_person) if default_person in options else 1
            score_hint = "、".join(f"{p}({s:.0%})" for p, s, _ in item["candidates"][:3])
            choice = st.selectbox(
                f"⚠️ {item['fname']}　－　候選：{score_hint}",
                options, index=default_idx, key=f"pending_{item['file']}",
            )
            st.session_state.pending_selection[item["file"]] = (
                None if choice == "— 略過此檔案 —" else choice
            )

    if scan["low_matches"]:
        with st.expander(f"❓ 低信心度（< {MEDIUM_CONFIDENCE_THRESHOLD:.0%}），"
                          f"請手動指定或忽略：{len(scan['low_matches'])} 個檔案", expanded=True):
            for item in scan["low_matches"]:
                options = ["— 不指定（略過）—"] + scan["names"]
                choice = st.selectbox(
                    item["fname"], options, index=0, key=f"low_{item['file']}",
                )
                st.session_state.low_selection[item["file"]] = (
                    None if choice.startswith("—") else choice
                )

    st.divider()
    process_btn = st.button("🚀 Step 2｜確認配對並開始處理", type="primary")

    if process_btn:
        run_finalize_and_store(scan)
else:
    st.info("請先於左側上傳人員名單 Excel 與證明文件 ZIP，並點擊「① 掃描並比對人員檔案」。")

tab1, tab2, tab3 = st.tabs(["📜 附錄證照效期檢核", "🧾 CV 結構化預覽與編輯", "⬇️ 下載"])

with tab1:
    if st.session_state.df_license is not None and not st.session_state.df_license.empty:
        st.dataframe(st.session_state.df_license, use_container_width=True, hide_index=True)
        status_series = st.session_state.df_license["狀態"].astype(str)
        n_expired = status_series.str.contains("已過期").sum()
        n_soon = status_series.str.contains("即將過期").sum()
        n_unknown = status_series.str.contains("⚠️").sum() - n_expired - n_soon
        if n_expired > 0:
            st.error(f"🔴 共有 {n_expired} 筆證照已過期，請優先處理！")
        if n_soon > 0:
            st.warning(f"🟡 共有 {n_soon} 筆證照即將於 90 天內到期。")
        if n_unknown > 0:
            st.info(f"⚠️ 共有 {n_unknown} 筆證照無法自動判讀效期，請人工確認。")
    elif st.session_state.df_license is not None:
        st.info("未偵測到任何技師執業執照或會員證檔案。")
    else:
        st.info("請先於左側上傳檔案並依序完成 Step 1 掃描配對與 Step 2 開始處理。")

with tab2:
    if st.session_state.df_staffing is not None:
        diagnostics = st.session_state.diagnostics or {}

        # ------------------------------------------------------------
        # ① Mapping 與離線萃取狀態儀表板
        # ------------------------------------------------------------
        total_people = len(diagnostics) if diagnostics else len(st.session_state.df_staffing)
        n_success = sum(1 for d in diagnostics.values() if d.get("status") == "🟢 離線萃取完整")
        n_warning = sum(1 for d in diagnostics.values() if d.get("status") == "🟡 資料待補")
        n_missing = sum(1 for d in diagnostics.values() if d.get("status") == "🔴 缺少 CV 檔案")
        n_unmatched = len(st.session_state.unmatched_files)

        st.markdown("#### 📊 Mapping 與離線萃取狀態儀表板")
        m1, m2, m3, m4, m5 = st.columns(5)
        m1.metric("應處理人數", total_people)
        m2.metric("🟢 CV 解析成功", n_success)
        m3.metric("🟡 資料待補（OCR等）", n_warning)
        m4.metric("🔴 缺少 CV 檔案", n_missing)
        m5.metric("⚠️ 未比對檔案數", n_unmatched)

        # ------------------------------------------------------------
        # ② 警示區塊
        # ------------------------------------------------------------
        missing_people = [name for name, d in diagnostics.items() if d.get("status") == "🔴 缺少 CV 檔案"]
        if missing_people:
            st.error(
                "以下人員缺乏 CV 履歷檔，已套用通用樣板保底（欄位不會空白，"
                "但內容較籠統），建議於下方表格手動補充精確資訊或重新上傳 CV：\n\n"
                + "、".join(missing_people)
            )

        warning_people = [(name, d.get("reason", "")) for name, d in diagnostics.items()
                          if d.get("status") == "🟡 資料待補"]
        if warning_people:
            lines = "\n".join(f"- **{name}**：{reason}" for name, reason in warning_people)
            st.warning("以下人員的 CV 內文擷取不完整（可能為純掃描影像，OCR 失敗），"
                       "已套用通用樣板保底，欄位皆有內容，建議檢查原因並視需要手動調整：\n\n" + lines)

        if st.session_state.unmatched_files:
            with st.expander(f"⚠️ 有 {len(st.session_state.unmatched_files)} 個檔案未成功比對到任何人員，"
                              "可於此手動指定並重新處理", expanded=False):
                st.caption("選擇對應人員後，點擊下方「套用修正並重新處理」即可將該檔案納入計算"
                           "（會重新執行 OCR 掃描，需再花費一些時間）。")
                for item in st.session_state.unmatched_files:
                    options = ["— 不指定（略過）—"] + st.session_state.scan["names"]
                    current = st.session_state.low_selection.get(item["file"])
                    default_idx = options.index(current) if current in options else 0
                    choice = st.selectbox(
                        item["fname"], options, index=default_idx,
                        key=f"tab2_fix_{item['file']}",
                    )
                    st.session_state.low_selection[item["file"]] = (
                        None if choice.startswith("—") else choice
                    )
                if st.button("🔄 套用修正並重新處理", key="tab2_apply_fix"):
                    run_finalize_and_store(st.session_state.scan)
                    st.rerun()

        # ------------------------------------------------------------
        # ③ 每位人員檔案配對明細（可摺疊）
        # ------------------------------------------------------------
        if diagnostics:
            with st.expander("🔍 檢視每位人員配對到的完整檔案清單（CV / 學歷 / 證照 / 投保）",
                              expanded=False):
                for name in st.session_state.scan["names"] if st.session_state.scan else diagnostics.keys():
                    diag = diagnostics.get(name)
                    if diag:
                        st.markdown(format_mapping_line(name, diag))

        st.divider()

        # ------------------------------------------------------------
        # ③.5 證照勾選 ＋ Summary 集中輸入縮寫面板
        # ------------------------------------------------------------
        license_state = st.session_state.license_state or {}
        if license_state:
            st.markdown("#### 🪪 證照/資格勾選與縮寫指定")
            st.caption("系統已從每位同仁 CV 的「資格」區塊精準擷取證照/資格全稱候選"
                       "（不會誤抓工作經歷段落），預設全部勾選。請於下方各人員面板勾選"
                       "是否採用；縮寫代碼改為在上方 Summary 看板「集中填寫一次」，"
                       "相同證照全稱只需輸入一次縮寫，全案自動套用。")

            person_order = st.session_state.scan["names"] if st.session_state.scan else list(license_state.keys())

            # 預留 Summary 看板版位（視覺上置頂），實際內容於下方蒐集完最新
            # 勾選狀態後才填入，確保 Summary 反映「這次互動後」的最新結果。
            summary_placeholder = st.container()

            # --- 每位人員的勾選編輯區（依 Excel 順序摺疊呈現，僅勾選框）---
            st.markdown("##### 👤 個人證照/資格勾選")
            for name in person_order:
                items = license_state.get(name, [])
                if not items:
                    continue
                checked_count = sum(1 for it in items if it.get("checked"))
                with st.expander(f"{name}（{len(items)} 項證照/資格，已勾選 {checked_count} 項）",
                                  expanded=False):
                    for idx, item in enumerate(items):
                        checked = st.checkbox(
                            item["text"], value=item.get("checked", True),
                            key=f"lic_check_{name}_{idx}",
                        )
                        item["checked"] = checked
            st.session_state.license_state = license_state

            # --- 依最新勾選狀態，彙整去重後的證照全稱清單 ---
            unique_licenses = get_unique_checked_licenses(license_state, person_order)
            abbrev_map = st.session_state.license_abbrev_map or {}

            # --- 回填頂部 Summary 看板：去重後每種證照全稱僅一個縮寫輸入框 ---
            with summary_placeholder:
                st.markdown("##### 📌 Summary：集中輸入縮寫（去重後，全案共用）")
                if not unique_licenses:
                    st.caption("目前沒有任何已勾選的證照/資格，請先於下方展開人員面板勾選。")
                else:
                    for lic_text in unique_licenses:
                        c1, c2 = st.columns([3, 1])
                        with c1:
                            st.write(lic_text)
                        with c2:
                            abbrev_val = st.text_input(
                                "縮寫", value=abbrev_map.get(lic_text, ""),
                                key=f"abbrev_global_{lic_text}",
                                label_visibility="collapsed", placeholder="縮寫",
                            )
                        abbrev_map[lic_text] = abbrev_val

                    stats = compute_abbrev_stats(license_state, abbrev_map)
                    if stats:
                        stats_line = "｜".join(
                            f"【{code}】{count}人" for code, count in
                            sorted(stats.items(), key=lambda x: -x[1])
                        )
                        st.info(f"📊 縮寫統計：{stats_line}")
                    else:
                        st.caption("尚未指定任何縮寫代碼，Badges 欄位目前皆為空白。")

            st.session_state.license_abbrev_map = abbrev_map

            # --- 即時連動：依目前勾選狀態 + 全局縮寫對照，重新計算 Licenses／Badges ---
            st.session_state.df_staffing = recompute_licenses_badges(
                st.session_state.df_staffing, license_state, abbrev_map
            )

        st.divider()

        # ------------------------------------------------------------
        # ④ 可編輯表格
        # ------------------------------------------------------------
        st.caption("可直接於表格中編輯任一欄位，或於最下方新增協力廠商（Subcontractor）資料列。"
                   "Licenses／Badges 欄位建議透過上方勾選面板編輯，於此表格中的修改會在下次"
                   "勾選面板互動時被覆蓋。")
        edited_df = st.data_editor(
            st.session_state.df_staffing,
            num_rows="dynamic",
            use_container_width=True,
            column_config={
                "Layer": st.column_config.SelectboxColumn(options=LAYER_OPTIONS),
                "YearsOfExp": st.column_config.NumberColumn(min_value=0, step=1),
            },
            key="staffing_editor",
        )
        st.session_state.df_staffing = edited_df
    else:
        st.info("請先於左側上傳檔案並依序完成 Step 1 掃描配對與 Step 2 開始處理。")

with tab3:
    if st.session_state.df_staffing is not None:
        col1, col2, col3 = st.columns(3)

        with col1:
            try:
                excel_bytes = build_template_excel(st.session_state.df_staffing)
                st.download_button(
                    "⬇️ 下載 Template_Staffing.xlsx",
                    data=excel_bytes,
                    file_name="Template_Staffing.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                )
            except Exception as e:
                st.error(f"Excel 產生失敗：{e}")

        with col2:
            if st.session_state.merged_pdf_bytes:
                st.download_button(
                    "⬇️ 下載 附錄_人員資格證明檔_Merged.pdf",
                    data=st.session_state.merged_pdf_bytes,
                    file_name="附錄_人員資格證明檔_Merged.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                )
            else:
                st.button("⬇️ 附錄 PDF（尚無資料）", disabled=True, use_container_width=True)

        with col3:
            try:
                license_excel = build_license_report_excel(
                    st.session_state.df_license if st.session_state.df_license is not None
                    else pd.DataFrame(columns=["姓名", "文件類別", "檔名", "狀態", "起始日", "截止日", "備註"])
                )
                st.download_button(
                    "⬇️ 下載 附錄人員資格與證照效期檢核報告.xlsx",
                    data=license_excel,
                    file_name="附錄人員資格與證照效期檢核報告.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                )
            except Exception as e:
                st.error(f"檢核報告產生失敗：{e}")
    else:
        st.info("請先於左側上傳檔案並依序完成 Step 1 掃描配對與 Step 2 開始處理。")

st.divider()
with st.expander("ℹ️ 使用說明與注意事項"):
    st.markdown(f"""
- **100% 本地離線運作**：本工具不呼叫任何外部 AI／雲端 API（不使用 Gemini 或
  任何其他 LLM 服務），也沒有任何 API 金鑰設定。所有 PDF 轉檔、OCR、文字
  萃取、證照效期判讀、Excel／PDF 產出，全部在你執行 `streamlit run app.py`
  的這台機器上完成，人員資料與標案內容不會外流。
- **人員名單 Excel** 必須包含欄位：`順序`、`姓名`、`部門`、`團隊職務`；
  建議另增 `英文姓名`（如 `Joe Lim`、`RAY HSU`）以提升檔名比對準確率，及
  `資歷`／`年資`／`YearsOfExp`／`工作年資` 其中一欄，作為找不到投保證明時的
  年資備援來源；未提供也不影響程式運作。
- **三段式雙軌配對機制**：以「姓名」與「英文姓名」為比對清單，比對前一律轉小寫、
  忽略大小寫，並移除 `_` `-` `()` `（）` 空格等符號差異，容許少量拼字誤差
  （如 Brain/Brian）：
  - 信心度 ≥ {HIGH_CONFIDENCE_THRESHOLD:.0%}：自動配對，於「已自動配對」清單顯示
    （如「已自動配對：Brian Lo → 羅翊軒」）。
  - 信心度 {MEDIUM_CONFIDENCE_THRESHOLD:.0%}～{HIGH_CONFIDENCE_THRESHOLD:.0%}：預設選中最高機率候選人，
    可於畫面上一鍵確認或修正。
  - 信心度 < {MEDIUM_CONFIDENCE_THRESHOLD:.0%}：需手動由下拉選單指定人員，或選擇略過。
- **內文導向（Content-Driven）分類**：
  - CV：檔名（不分大小寫）含「CV／履歷／學經歷／經歷表」，或副檔名為
    `.doc`／`.docx`（含大寫 `.DOC`）→ 直接判定為 CV，不需 OCR。
  - 技師執業執照：需先對 PDF／圖片進行文字擷取（含 OCR 備援），**內文**包含
    「技師執業執照」或「執業執照」才成立，不再依賴檔名。
  - 技師公會會員證：**內文同時**包含「會員證」與「技師公會」（如：台北市水利
    技師公會、台灣省水利技師公會）才成立，避免景觀學會等一般協會會員證被
    誤判。
  - 投保證明／學歷證明：內容結構化程度低，維持以檔名關鍵字判斷。
  - 其餘（一般協會會員證、電子證書等）：一律歸為一般證照，不進行過期告警。
- **內文與 OCR 掃描進度**：Step 2 開始處理時，會先以進度條與狀態訊息顯示
  「正在進行內文與 OCR 掃描：[檔名] (X/Y)」，掃描完成後才進入離線欄位萃取階段。
- **離線欄位萃取（零 API 消耗、零失敗）**：CV 全文一律完整讀取，不做任何頁數
  ／字數截斷；以下欄位純用 Python Regex／規則式邏輯離線判定，永遠有內容：
  - `Degree`：依「博士＞碩士＞研究所＞學士＞大學」優先序，在 CV 中搜尋含學歷
    關鍵字的段落，並強制清除「學歷：」「學 歷：」等字頭（含全／半形空白與
    冒號的各種組合）。
  - `Company`／`Title`：辨識「現職／現 職／現任／服務單位」等表頭（相容
    `現    職：` 這類夾帶多個半形/全形空白或 Tab 的寫法）；表頭與內容分行時
    （如「現職：」單獨一行，下一行才是「-美商傑明工程顧問(股)台灣分公司
    工程師」）會自動向下掃描最多 3 行、去除開頭的「-」項目符號後再精準切分
    公司全名與職稱；若切分未抓到職稱，改用常用職稱字典（正工程師／水利
    工程師／監造主任／資深協理／工程師…）掃描 CV 前 15 行補全。全程禁止把
    「職稱」「服務單位」等表頭字樣本身填入欄位。
  - `Licenses`／`Badges`：**已移除所有自動縮寫字典與 Hardcode 規則**。系統僅
    離線擷取每位同仁「證照/資格全稱候選清單」（來源：CV 中學歷／證照／技師／
    資格相關段落，以及已分類為執業執照／會員證的檔名），完全不自動判定縮寫。
    實際輸出的 `Licenses`（全稱，頓號分隔）與 `Badges`（縮寫，逗號分隔）由
    你在 tab2「🪪 證照/資格勾選與縮寫指定」面板中，逐項勾選是否採用、並手動
    輸入縮寫代碼後即時運算產生（詳見下方說明）。
  - `Expertise`：優先擷取 CV「專長」段落，清除項目符號／控制字元等亂碼後，
    僅保留 2~10 字的精簡詞彙，統一輸出 4~6 個以「/」分隔的關鍵字（過濾長句
    敘述）；找不到則以已擷取到的證照/資格候選清單（前幾筆）展開為專長描述。
  - `YearsOfExp`：優先由 `6_投保證明`「勞保投保年資：X年 X日」離線換算
    （日數 ≥180 進位 +1年，如 6年309日→7年）；找不到投保證明時，退而讀取
    人員名單 Excel 中的資歷欄位（`資歷`／`年資`／`YearsOfExp`／`工作年資`，
    偵測到即用）；兩者皆無資料則預設為 0。
  - **CV 主要工作/職務內容**：從 CV「經歷／工作內容／專案職責」等段落，擷取
    2~3 句含「負責/辦理/執行/規劃…」等動詞的實際工作內容，供 `JobDescription`
    離線生成使用。
  - `JobDescription`：依團隊職務（Role）、部門/組別（GroupName）**與 CV 主要
    工作內容**三者動態組合（如「於細部設計組擔任組員，負責降漏水率管網規劃
    、水理模式建置與現場勘查」），不會所有人套用同一句話；找不到 CV 工作
    內容時，退回角色/組別對應樣板。
  - `BioNarrative`：採四段式標案範本風格組成完整敘述——①學歷與現職
    ②代表性經歷（CV 中 1~2 項重點專案名稱） ③核心專長 ④履約效益（團隊
    協調與品質控管能力）。
  - 以上 `BioNarrative`／`JobDescription` 內容已盡量豐富完整，可直接作為
    最終輸出使用；若想要更精緻的文字潤飾，也可以自行將 Excel 內容複製到
    Gemini 網頁版或其他工具，由你手動調整語氣與用字。
- **🪪 證照/資格勾選 ＋ Summary 集中輸入縮寫面板**（tab2，位於配對明細與可編輯
  表格之間）：
  - 證照/資格候選改用**區塊邊界精準擷取**：僅讀取 CV 中「資格：」「資 格：」
    「專業資格：」「證照：」「資格及訓練：」等標頭，一路擷取到下一個「教育
    訓練：」「工作經歷：」「經歷：」「專案經歷：」標頭為止；找不到明確區塊
    才退回較寬鬆的關鍵字比對。嚴禁掃描工作經歷段落，避免「擔任設計技師」
    「監造技師」等專案說明文字被誤抓為證照。
  - 每位同仁以 `st.expander` 摺疊呈現其擷取到的證照/資格全稱候選，**僅提供
    勾選框**（預設全部勾選），不再逐項填縮寫。
  - 縮寫代碼改為**頂部 Summary 看板集中輸入**：自動彙整全案「目前已勾選」的
    證照全稱並去重複，去重後的每種全稱只出現一個縮寫輸入框——多人持有同一
    張證照，縮寫只需填寫一次、全案自動套用；旁邊即時顯示各縮寫代碼的統計
    （如「【技】4人｜【品】3人｜【P】2人」）。
  - 即時連動：任何勾選/取消勾選或修改縮寫的操作，都會立即重新計算並寫回
    `Licenses`（該同仁已勾選項目的全稱，頓號分隔）與 `Badges`（該同仁已勾選
    且其全稱已被填寫縮寫的代碼，半形逗號分隔，同一代碼不重複列出）兩欄，
    下方表格與下載的 Excel 皆同步反映最新結果。手動於表格新增的列（如協力
    廠商）不受此面板影響，其 `Licenses`／`Badges` 由你直接在表格中填寫即可。
- **`determine_layer` 異體字相容**：「計畫」與「計劃」（畫／劃異體字）視為
  同義，「計畫主持人」「計劃主持人」「副計畫主持人」「專案經理」等寫法皆可
  正確判定 Layer，避免因用字差異被誤判為 `GroupMember`。
- **CV 文字擷取**：PDF（含大寫 `.PDF`，內建 OCR 備援）與 Word（含大寫
  `.DOC`／`.DOCX`）皆會完整擷取全文；輸出的 Template_Staffing.xlsx 為新版
  14 欄格式（較原標準格式新增 `Licenses` 欄，位於 `Company` 與 `Badges`
  之間）。
- **合併 PDF 雙層嚴格排序**：第一層依人員名單 Excel「順序」欄由上到下排列
  人員；第二層每人內部嚴格依 `1_CV → 2_學歷 → 3_證照 → 4_技師執業執照 →
  5_技師會員證 → 6_投保證明` 順序合併，並加入書籤方便導覽。
- **證照效期檢核（逐頁獨立解析）**：僅針對「技師執業執照」與「技師會員證」
  兩類進行，其餘證照類別不判斷過期。**採逐頁（Page-by-Page）OCR 文字擷取
  與獨立檢核**，可正確處理如「潘冠愷合併.pdf」這類將多張證照掃描影像合併
  於同一份 PDF 的情況（例如第2頁是執業執照、第3頁是會員證），不會因整份
  檔案只被歸類為單一種類而漏掉其他頁面的證照效期。日期格式支援「執照有效
  期間：自民國X年X月X日至X年X月X日止」（可跨行/任意空格，如「自民國113年
  \n1月31日至119年1月30日止」）、「有效期限：民國X年X月X日至X年X月X日」、
  「115年會員證／115年度會員證」（自動推算至該年12月31日，如
  115+1911=2026年12月31日），以及**專門抓取「至／~／止」後方截止日的備援
  規則**（因應 OCR 誤判或格式跑版導致起訖日全段比對失敗）；日期一律以當下
  系統日期（`date.today()`）動態計算。**證照與投保證明的解析全程為離線
  Regex／OCR，不呼叫任何外部 API。**
- **表格編輯即時同步**：`tab2` 的 `st.data_editor` 編輯結果會立即寫回
  `st.session_state.df_staffing`；`tab3` 下載 `Template_Staffing.xlsx` 時
  一律讀取當下最新的 `st.session_state.df_staffing`，因此任何手動修改
  （Layer／BioNarrative 等）都會反映在下載檔案中。
- 部署到 **Streamlit Community Cloud** 時，請將本工具一併產生的 `packages.txt`
  放在 repo 根目錄，以安裝 LibreOffice 與 Tesseract 等系統套件；不需要設定
  任何 Secrets 或 API 金鑰。
""")
